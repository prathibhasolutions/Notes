"""
PPT to Word Textbook Converter — Local Server
Run:  python ppt_to_word_server.py
Open: http://localhost:5000
"""

import io
import os
import re
import traceback
import xml.etree.ElementTree as ET
import zipfile

from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
CORS(app)


# ── Document styles ────────────────────────────────────────────────────────────

def setup_styles(doc):
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    s = doc.styles['Normal']
    s.font.name = 'Calibri'
    s.font.size = Pt(11)

    s = doc.styles['Heading 1']
    s.font.name = 'Calibri'
    s.font.size = Pt(17)
    s.font.bold = True
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        s.font.color.rgb = RGBColor(0x17, 0x29, 0x54)
    except Exception:
        pass

    s = doc.styles['Heading 2']
    s.font.name = 'Calibri'
    s.font.size = Pt(13)
    s.font.bold = True
    try:
        s.font.color.rgb = RGBColor(0x1E, 0x3A, 0x6E)
    except Exception:
        pass


# ── Paragraph spacing helper ──────────────────────────────────────────────────

def set_spacing(para, before=0, after=3, line=276):
    """
    Set paragraph spacing directly on the XML element.
    line is in twips: 240 = single, 276 ≈ 1.15x, 360 = 1.5x
    before/after are in points (converted to twips internally).
    """
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'),    str(int(before * 20)))
    sp.set(qn('w:after'),     str(int(after  * 20)))
    sp.set(qn('w:line'),      str(int(line)))
    sp.set(qn('w:lineRule'),  'auto')
    pPr.append(sp)


# ── Font-size mapper ───────────────────────────────────────────────────────────

def textbook_pt(ppt_pt):
    """
    Map oversized PPT font sizes down to compact textbook sizes.
    PPT typically uses 24–40 pt; we want 10–13 pt in Word.
    """
    if ppt_pt is None:
        return 11
    if ppt_pt >= 40:
        return 13
    if ppt_pt >= 28:
        return 12
    if ppt_pt >= 20:
        return 11
    return 10


# ── Decorative bottom border under a paragraph ───────────────────────────────

def _add_bottom_border(para, color='172954', sz=8, space=4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    str(sz))
    bottom.set(qn('w:space'), str(space))
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Write styled runs into a Word paragraph ───────────────────────────────────

def _write_runs(para, runs, default_size=11):
    """Render a list of {text, bold, italic, underline, size} run dicts into para."""
    for run_data in runs:
        r = para.add_run(run_data['text'])
        r.bold   = run_data.get('bold',   False)
        r.italic = run_data.get('italic', False)
        if run_data.get('underline', False):
            r.underline = True
        size        = run_data.get('size')
        r.font.size = Pt(textbook_pt(size) if size else default_size)
        r.font.name = 'Calibri'


# ── XML namespace helpers ─────────────────────────────────────────────────────
# PPTX is a plain ZIP file.  We read slide XML with the standard library and
# never create python-pptx shape objects, so 'shape is not a placeholder'
# cannot be raised under any circumstances.

_NS = {
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'p': 'http://schemas.openxmlformats.org/presentationml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}

def _q(tag):
    """'a:t'  →  '{http://...}t'  (Clark notation for ElementTree)"""
    prefix, local = tag.split(':', 1)
    return '{%s}%s' % (_NS[prefix], local)


def _parse_title(root):
    """Return slide title text from an ET root element, or ''."""
    try:
        for sp in root.iter(_q('p:sp')):
            ph = sp.find('.//' + _q('p:ph'))
            if ph is None:
                continue
            if ph.get('type', '') in ('title', 'ctrTitle') or ph.get('idx', '99') == '0':
                text = ''.join(t.text for t in sp.iter(_q('a:t')) if t.text).strip()
                if text:
                    return text
    except Exception:
        pass
    return ''


def _is_ctr_title_slide(root):
    """Return True if slide has a ctrTitle placeholder (title/section-divider slide)."""
    try:
        for sp in root.iter(_q('p:sp')):
            ph = sp.find('.//' + _q('p:ph'))
            if ph is not None and ph.get('type', '') == 'ctrTitle':
                return True
    except Exception:
        pass
    return False


def _parse_body(root):
    """
    Return body paragraphs as list of dicts:
      { text, runs:[{text, bold, italic, size}], level, bold, size, has_bullet }
    has_bullet is True unless the paragraph has an explicit <a:buNone/> element.
    """
    result = []
    try:
        for sp in root.iter(_q('p:sp')):
            ph = sp.find('.//' + _q('p:ph'))
            if ph is not None:
                if ph.get('type', '') in ('title', 'ctrTitle') or ph.get('idx', '99') == '0':
                    continue
            txBody = sp.find(_q('p:txBody'))
            if txBody is None:
                continue
            for p_elem in txBody.findall(_q('a:p')):
                level      = 0
                has_bullet = True   # PPT body text is bulleted by default
                align      = 'left'
                pPr = p_elem.find(_q('a:pPr'))
                if pPr is not None:
                    try:
                        level = int(pPr.get('lvl', '0'))
                    except Exception:
                        level = 0
                    if pPr.find(_q('a:buNone')) is not None:
                        has_bullet = False
                    algn_raw = pPr.get('algn', '')
                    align = {'ctr': 'center', 'r': 'right',
                             'just': 'justify', 'dist': 'justify'}.get(algn_raw, 'left')

                runs = []
                for r_elem in p_elem.findall(_q('a:r')):
                    t_elem = r_elem.find(_q('a:t'))
                    if t_elem is None or not t_elem.text:
                        continue
                    bold      = False
                    italic    = False
                    underline = False
                    size      = None
                    rPr = r_elem.find(_q('a:rPr'))
                    if rPr is not None:
                        bold      = rPr.get('b', '0') not in ('0', 'false')
                        italic    = rPr.get('i', '0') not in ('0', 'false')
                        underline = rPr.get('u', 'none') not in ('none', '')
                        sz        = rPr.get('sz')
                        if sz:
                            try:
                                size = int(sz) / 100.0
                            except Exception:
                                pass
                    runs.append({'text': t_elem.text, 'bold': bold,
                                 'italic': italic, 'underline': underline, 'size': size})

                if not runs:
                    continue
                full_text  = ''.join(r['text'] for r in runs).strip()
                if not full_text:
                    continue
                all_bold   = all(r['bold'] for r in runs)
                first_size = next((r['size'] for r in runs if r['size']), None)
                result.append({
                    'text':       full_text,
                    'runs':       runs,
                    'level':      level,
                    'bold':       all_bold,
                    'size':       first_size,
                    'has_bullet': has_bullet,
                    'align':      align,
                })
    except Exception:
        pass
    return result


def _parse_tables(root):
    """Return tables as list-of-list-of-str from an ET root element."""
    tables = []
    try:
        for tbl in root.iter(_q('a:tbl')):
            try:
                rows = []
                for tr in tbl.findall(_q('a:tr')):
                    row = [''.join(t.text for t in tc.iter(_q('a:t')) if t.text).strip()
                           for tc in tr.findall(_q('a:tc'))]
                    if row:
                        rows.append(row)
                if rows:
                    tables.append(rows)
            except Exception:
                continue
    except Exception:
        pass
    return tables


def _parse_images(root, rels_root, zf, names):
    """
    Return list of {blob, width_in} for each image on the slide.
    Width is read from the PPTX XML <a:ext cx="..."> attribute (EMU → inches),
    capped at MAX_W so images never overflow the page.
    """
    result  = []
    rid_map = {}
    if rels_root is not None:
        for rel in rels_root:
            if 'image' in rel.get('Type', '').lower():
                target   = rel.get('Target', '')
                img_path = ('ppt/' + target[3:]) if target.startswith('../') else target
                rid_map[rel.get('Id', '')] = img_path

    R_EMBED = '{%s}embed' % _NS['r']
    EMU     = 914400.0   # English Metric Units per inch
    MAX_W   = 5.5        # max image width in inches (fits within 1.25" margins)

    for pic in root.iter(_q('p:pic')):
        try:
            blip = pic.find('.//' + _q('a:blip'))
            if blip is None:
                continue
            r_id     = blip.get(R_EMBED)
            img_path = rid_map.get(r_id or '', '')
            if not img_path or img_path not in names:
                continue
            blob = zf.read(img_path)

            # Derive display width from the shape's bounding box in the slide XML
            w_in = 4.5  # sensible fallback
            ext  = pic.find('.//' + _q('a:ext'))
            if ext is not None:
                try:
                    cx = int(ext.get('cx', 0))
                    if cx > 0:
                        w_in = min(cx / EMU, MAX_W)
                except Exception:
                    pass

            result.append({'blob': blob, 'width': w_in})
        except Exception:
            continue
    return result


# ── Core conversion ────────────────────────────────────────────────────────────

def pptx_to_docx(pptx_bytes: bytes) -> bytes:
    """
    Convert a .pptx (bytes) to textbook-style .docx (bytes).
    Opens PPTX as a plain ZIP — zero python-pptx shape objects created,
    so 'shape is not a placeholder' is structurally impossible.
    Slide backgrounds are skipped entirely (not needed in Word).
    """
    doc = Document()
    setup_styles(doc)
    try:
        ep = doc.paragraphs[0]._element
        ep.getparent().remove(ep)
    except Exception:
        pass

    try:
        zf = zipfile.ZipFile(io.BytesIO(pptx_bytes))
    except Exception as exc:
        raise ValueError(f'Could not open as PowerPoint file: {exc}')

    is_first_slide = True

    with zf:
        names = set(zf.namelist())

        # Slide files sorted numerically: slide1.xml, slide2.xml, ...
        slide_files = sorted(
            [n for n in names if re.match(r'ppt/slides/slide\d+\.xml$', n)],
            key=lambda n: int(re.search(r'\d+', n.split('/')[-1]).group())
        )

        for slide_file in slide_files:
            try:
                root = ET.fromstring(zf.read(slide_file))

                title_text = _parse_title(root)
                body_paras = _parse_body(root)
                tables     = _parse_tables(root)
                is_ctr     = _is_ctr_title_slide(root)

                # Parse images with original size info from the slide XML
                rels_root = None
                rels_path = slide_file.replace('ppt/slides/', 'ppt/slides/_rels/') + '.rels'
                if rels_path in names:
                    try:
                        rels_root = ET.fromstring(zf.read(rels_path))
                    except Exception:
                        pass
                images = _parse_images(root, rels_root, zf, names)

                if not title_text and not body_paras and not tables and not images:
                    continue  # blank slide — skip

                # Classify slide type for smart rendering:
                #   is_ctr      → title / section-divider slide (ctrTitle placeholder)
                #   is_diagram  → image(s) with no body text (diagram / figure slide)
                #   else        → normal content slide
                is_diagram = bool(images) and not bool(body_paras) and not bool(tables)

                # ── Section separator between slides ───────────────────────────
                if not is_first_slide:
                    sep = doc.add_paragraph('─' * 72)
                    try:
                        sep.runs[0].font.size = Pt(7)
                        sep.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
                    except Exception:
                        pass
                    set_spacing(sep, before=8, after=8, line=240)
                is_first_slide = False

                _ALIGN_MAP = {
                    'center':  WD_ALIGN_PARAGRAPH.CENTER,
                    'right':   WD_ALIGN_PARAGRAPH.RIGHT,
                    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
                    'left':    WD_ALIGN_PARAGRAPH.LEFT,
                }

                # ═════════════════════════════════════════════════════
                # PATH A: TITLE / SECTION-DIVIDER SLIDE
                # Prominent centered title + subtitle text in italic
                # ═════════════════════════════════════════════════════
                if is_ctr:
                    if title_text:
                        h = doc.add_heading(title_text, level=1)
                        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        _add_bottom_border(h, color='172954', sz=8, space=4)
                        set_spacing(h, before=14, after=8)
                    for para in body_paras:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        set_spacing(p, before=2, after=3)
                        for rd in para['runs']:
                            r = p.add_run(rd['text'])
                            r.italic    = True
                            r.font.size = Pt(12)
                            r.font.name = 'Calibri'

                # ═════════════════════════════════════════════════════
                # PATH B: DIAGRAM / FIGURE SLIDE
                # Image(s) first, centered + proportional;
                # slide title rendered as a centered italic figure caption below.
                # ═════════════════════════════════════════════════════
                elif is_diagram:
                    for img in images:
                        try:
                            doc.add_picture(io.BytesIO(img['blob']), width=Inches(img['width']))
                            pic_para = doc.paragraphs[-1]
                            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            set_spacing(pic_para, before=4, after=2, line=240)
                        except Exception:
                            continue
                    if title_text:
                        cap = doc.add_paragraph()
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = cap.add_run(title_text)
                        r.italic    = True
                        r.bold      = False
                        r.font.size = Pt(10)
                        r.font.name = 'Calibri'
                        try:
                            r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                        except Exception:
                            pass
                        set_spacing(cap, before=2, after=10)

                # ═════════════════════════════════════════════════════
                # PATH C: CONTENT SLIDE
                # Title as heading → body bullets/paragraphs → tables → images
                # Paragraph alignment respected from PPT XML.
                # ═════════════════════════════════════════════════════
                else:
                    if title_text:
                        h = doc.add_heading(title_text, level=1)
                        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        _add_bottom_border(h, color='172954', sz=6, space=3)
                        set_spacing(h, before=10, after=8)

                    for para in body_paras:
                        try:
                            text       = para['text']
                            level      = para['level']
                            all_bold   = para['bold']
                            size_pt    = para['size']
                            has_bullet = para['has_bullet']
                            runs       = para['runs']
                            align      = para.get('align', 'left')

                            # Sub-heading: all-bold with large PPT font
                            if all_bold and size_pt and size_pt >= 20:
                                p = doc.add_heading(text, level=2)
                                set_spacing(p, before=6, after=3)

                            # Bullet / indented item
                            elif level > 0 or has_bullet:
                                p = doc.add_paragraph(style='List Bullet')
                                if level > 0:
                                    p.paragraph_format.left_indent = Inches(0.25 * level)
                                p.alignment = _ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.LEFT)
                                set_spacing(p, before=0, after=2, line=276)
                                _write_runs(p, runs)

                            else:
                                # Plain paragraph (explicit buNone in PPT)
                                p = doc.add_paragraph()
                                p.alignment = _ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.LEFT)
                                set_spacing(p, before=2, after=3)
                                _write_runs(p, runs)

                        except Exception:
                            continue

                    for table_data in tables:
                        try:
                            rows = len(table_data)
                            cols = max(len(row) for row in table_data)
                            if rows < 1 or cols < 1:
                                continue
                            wt = doc.add_table(rows=rows, cols=cols)
                            wt.style = 'Table Grid'
                            for ri, row_data in enumerate(table_data):
                                for ci, cell_text in enumerate(row_data):
                                    try:
                                        cell = wt.cell(ri, ci)
                                        cell.text = cell_text
                                        for run in cell.paragraphs[0].runs:
                                            run.font.name = 'Calibri'
                                            run.font.size = Pt(10)
                                            if ri == 0:
                                                run.bold = True  # bold header row
                                    except Exception:
                                        pass
                            sp = doc.add_paragraph()
                            set_spacing(sp, before=0, after=6, line=240)
                        except Exception:
                            continue

                    for img in images:
                        try:
                            doc.add_picture(io.BytesIO(img['blob']), width=Inches(img['width']))
                            pic_para = doc.paragraphs[-1]
                            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            set_spacing(pic_para, before=6, after=6, line=240)
                        except Exception:
                            continue  # skip unsupported formats (EMF, WMF, etc.)

            except Exception:
                continue  # skip any slide that fails for any reason

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


# ── Flask routes ───────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return send_file('ppt_to_word.html')


@app.route('/convert', methods=['POST'])
def do_convert():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    f = request.files['file']
    if not f.filename.lower().endswith('.pptx'):
        return jsonify({
            'error': 'Please upload a .pptx file (PowerPoint 2007 or later).'
        }), 400

    try:
        docx_bytes = pptx_to_docx(f.read())
        base_name  = os.path.splitext(f.filename)[0]
        out_name   = base_name + '_textbook.docx'
        return send_file(
            io.BytesIO(docx_bytes),
            as_attachment=True,
            download_name=out_name,
            mimetype=(
                'application/vnd.openxmlformats-officedocument'
                '.wordprocessingml.document'
            )
        )
    except Exception as exc:
        tb = traceback.format_exc()
        print('=== CONVERSION ERROR ===\n' + tb)
        return jsonify({'error': f'Conversion failed: {exc}', 'traceback': tb}), 500


# ── Entry point ────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print('PPT to Word Converter running at http://localhost:5000')
    print('Open the above URL in your browser.')
    app.run(debug=False, port=5000)
